#依赖包:
#pip install PyMuPDF python-docx pillow

import os
import io
import fitz
from docx import Document
from docx.shared import Inches
from PIL import Image

def pdf_to_word(pdf_file, word_file):
    doc = Document()
    pdf_document = fitz.open(pdf_file)

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        
        text = page.get_text()

        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image = Image.open(io.BytesIO(base_image["image"]))
            image_path = f"temp_image_{page_num}_{img_index}.png"
            image.save(image_path)

            doc.add_picture(image_path, width=Inches(6))

            os.remove(image_path)

        doc.add_paragraph(text)
        doc.add_page_break()
    
    doc.save(word_file)

def main():
    pdf_file = input("请输入要转换的PDF文件路径: ")
    output_directory = input("请输入输出目录（按回车使用默认目录）: ")

    if not output_directory:
        output_directory = os.path.dirname(pdf_file)

    output_file = os.path.join(output_directory, f"{os.path.splitext(os.path.basename(pdf_file))[0]}.docx")

    try:
        pdf_to_word(pdf_file, output_file)
        print("转换完成！")
    except Exception as e:
        print(f"转换失败：{str(e)}")

if __name__ == "__main__":
    main()

