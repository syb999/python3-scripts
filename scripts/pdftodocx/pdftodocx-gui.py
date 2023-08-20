#依赖包:
#pip install PySimpleGUI PyMuPDF python-docx pillow

import PySimpleGUI as sg
import os
import io
import fitz
from docx import Document
from docx.shared import Inches
from PIL import Image

def pdf_to_word(pdf_file, word_file):
    doc = Document()
    pdf_document = fitz.open(pdf_file)

    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        
        text = page.get_text()

        image_list = page.get_images(full=True)
        
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image = Image.open(io.BytesIO(base_image["image"]))
            image_path = f"temp_image_{page_num}_{img_index}.png"
            image.save(image_path)

            doc.add_picture(image_path, width=Inches(6))

            os.remove(image_path)

        doc.add_paragraph(text)
        doc.add_page_break()
    
    doc.save(word_file)

def main():
    sg.theme('LightGrey1')

    layout = [
        [sg.Text("选择输入的PDF文件:", justification='right', size=(15, 1)), sg.InputText(key="pdf_file", size=(50, 1)), sg.FileBrowse()],
        [sg.Text("更改输出目录:", justification='right', size=(15, 1)), sg.InputText(key="output_directory", size=(50, 1)), sg.FolderBrowse()],
        [sg.Button("转换"), sg.Button("退出")]
    ]

    window = sg.Window("PDF转Word工具", layout)

    while True:
        event, values = window.read()

        if event == sg.WIN_CLOSED or event == "退出":
            break
        elif event == "转换":
            pdf_file = values["pdf_file"]
            output_directory = values["output_directory"] or os.path.dirname(pdf_file)
            output_file = os.path.join(output_directory, f"{os.path.splitext(os.path.basename(pdf_file))[0]}.docx")

            try:
                pdf_to_word(pdf_file, output_file)
                sg.popup("转换完成！", title="提示")
            except Exception as e:
                sg.popup(f"转换失败：{str(e)}", title="错误提示")

    window.close()

if __name__ == "__main__":
    main()

